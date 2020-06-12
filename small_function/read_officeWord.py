# coding: utf-8
#! python3
# require python-docx

from docx import Document


f_path = 'H:/DF_emotion/标数据/modify.docx'
document = Document(f_path)
print(len(document.paragraphs))  # 获得行数
for paragraph in document.paragraphs:
    print(paragraph.text)  # 输出每一行
